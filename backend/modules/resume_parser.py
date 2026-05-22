"""
Resume Parser Module — Structured JSON Extraction with Gemini
Extracts fully structured information from resume files including
project-skill knowledge graph for resume-grounded question generation.
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

logger = logging.getLogger(__name__)

# ─── Gemini client (lazy import) ───────────────────────────────────────────────
def _get_gemini_model():
    try:
        import google.generativeai as genai
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return None
        genai.configure(api_key=api_key)
        return genai.GenerativeModel("gemini-1.5-flash")
    except Exception:
        return None


# ─── Structured JSON schema prompt ─────────────────────────────────────────────
_EXTRACTION_PROMPT = """
You are a precise resume parser. Extract ALL technical information from the resume text below and return ONLY a valid JSON object — no markdown, no explanation.

JSON Schema (fill every field; use empty list/dict if not present):
{{
  "candidate_name": "string",
  "email": "string",
  "phone": "string",
  "experience_years": integer,
  "skills": ["list of ALL technical skills, languages, frameworks, databases, cloud services"],
  "tools": ["list of specific tools, libraries, SDKs, APIs explicitly mentioned"],
  "certifications": ["list of certifications or courses"],
  "education": ["list of degrees e.g. B.Tech in Computer Science"],
  "experience": [
    {{
      "company": "string",
      "role": "string",
      "duration": "string (e.g. Jun 2023 - Dec 2023)",
      "skills": ["skills/technologies used in this role"]
    }}
  ],
  "projects": {{
    "<Project Name>": {{
      "description": "1-2 sentence description of what the project does",
      "skills": ["technologies / frameworks used"],
      "tools": ["specific tools, APIs, libraries used"],
      "outcome": "brief result or achievement if mentioned"
    }}
  }}
}}

IMPORTANT RULES:
- Extract ONLY what is explicitly stated in the resume
- Do NOT invent or infer technologies not mentioned
- For projects: list only technologies directly connected to THAT specific project
- Ignore leadership, NSS, extracurricular, hobbies sections entirely
- experience_years: count from earliest work/internship start year to present; 0 if no experience

Resume Text:
{resume_text}
"""


class ResumeParser:
    """Parse and extract structured information from resumes using Gemini."""

    def __init__(self):
        # Fallback skill keywords for regex path
        self.skill_keywords = {
            "Backend": ["python", "java", "node.js", "fastapi", "django", "spring",
                        "docker", "kubernetes", "flask", "go", "rust", "c#"],
            "AI/ML": ["python", "tensorflow", "pytorch", "scikit-learn", "nlp",
                      "deep learning", "machine learning", "keras", "opencv",
                      "huggingface", "langchain", "transformers"],
            "Frontend": ["react", "vue", "angular", "javascript", "typescript",
                         "html", "css", "next.js", "svelte"],
            "DevOps": ["docker", "kubernetes", "jenkins", "aws", "gcp",
                       "terraform", "ci/cd", "ansible", "github actions"],
            "Data Science": ["python", "pandas", "numpy", "sql", "hadoop",
                             "spark", "tableau", "power bi", "r"],
            "Tools": ["git", "github", "postman", "vscode", "jupyter",
                      "streamlit", "firebase", "sqlite", "postgresql",
                      "mysql", "mongodb", "redis", "gemini", "openai"],
        }
        self.domain_keywords = {
            "Finance": ["fintech", "trading", "banking", "payment", "risk"],
            "Healthcare": ["healthcare", "medical", "hospital", "pharma"],
            "E-Commerce": ["ecommerce", "retail", "marketplace", "shopping"],
            "SaaS": ["saas", "subscription", "cloud", "api"],
            "Data": ["analytics", "big data", "data engineering", "warehouse"],
        }

    # ─── Public API ────────────────────────────────────────────────────────────

    def parse(self, file_content: bytes, filename: str) -> Dict:
        """
        Parse resume file and return a fully structured data dictionary.

        Returns:
            Dict with keys: candidate_name, email, phone, experience_years,
            skills, tools, certifications, education, experience, projects,
            knowledge_graph, domain, raw_text
        """
        try:
            raw_text = self._read_file(file_content, filename)
        except Exception as e:
            logger.error(f"File extraction error: {e}")
            raise Exception(f"Resume file could not be read: {e}")

        # 1. Try Gemini structured extraction
        structured = self._gemini_extract(raw_text)

        # 2. Fallback to regex if Gemini unavailable
        if not structured:
            logger.warning("Gemini unavailable — using regex fallback parser")
            structured = self._regex_extract(raw_text)

        # 3. Always enrich with domain detection + knowledge graph
        structured["domain"] = self._detect_domain(raw_text)
        structured["knowledge_graph"] = self._build_knowledge_graph(structured)
        structured["raw_text"] = raw_text

        logger.info(
            f"Resume parsed: {structured.get('candidate_name')} | "
            f"{len(structured.get('skills', []))} skills | "
            f"{len(structured.get('projects', {}))} projects"
        )
        return structured

    # ─── Gemini Extraction ─────────────────────────────────────────────────────

    def _gemini_extract(self, raw_text: str) -> Optional[Dict]:
        """Use Gemini to extract structured JSON from resume text."""
        model = _get_gemini_model()
        if not model:
            return None

        try:
            prompt = _EXTRACTION_PROMPT.format(resume_text=raw_text[:8000])
            response = model.generate_content(prompt)
            text = response.text.strip()

            # Strip markdown code fences if present
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)

            data = json.loads(text)
            return self._validate_and_normalize(data)

        except json.JSONDecodeError as e:
            logger.warning(f"Gemini returned invalid JSON: {e}")
            return None
        except Exception as e:
            logger.error(f"Gemini extraction failed: {e}")
            return None

    def _validate_and_normalize(self, data: Dict) -> Dict:
        """Ensure all required fields exist with correct types."""
        result = {
            "candidate_name": str(data.get("candidate_name") or "Unknown"),
            "email": str(data.get("email") or "N/A"),
            "phone": str(data.get("phone") or "N/A"),
            "experience_years": int(data.get("experience_years") or 0),
            "skills": list(data.get("skills") or []),
            "tools": list(data.get("tools") or []),
            "certifications": list(data.get("certifications") or []),
            "education": list(data.get("education") or []),
            "experience": list(data.get("experience") or []),
            "projects": dict(data.get("projects") or {}),
        }
        # Normalize project values
        normalized_projects = {}
        for name, details in result["projects"].items():
            if isinstance(details, dict):
                normalized_projects[name] = {
                    "description": str(details.get("description") or ""),
                    "skills": list(details.get("skills") or []),
                    "tools": list(details.get("tools") or []),
                    "outcome": str(details.get("outcome") or ""),
                }
            else:
                normalized_projects[name] = {
                    "description": str(details),
                    "skills": [],
                    "tools": [],
                    "outcome": "",
                }
        result["projects"] = normalized_projects
        return result

    # ─── Regex Fallback ────────────────────────────────────────────────────────

    def _regex_extract(self, text: str) -> Dict:
        """Fallback regex-based extractor (original logic, extended)."""
        return {
            "candidate_name": self._extract_name(text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "experience_years": self._extract_experience_years(text),
            "skills": self._extract_skills_flat(text),
            "tools": [],
            "certifications": self._extract_certifications(text),
            "education": self._extract_education(text),
            "experience": [],
            "projects": self._extract_projects_dict(text),
        }

    def _extract_name(self, text: str) -> str:
        for line in text.split("\n")[:6]:
            line = line.strip()
            if line and 2 <= len(line.split()) <= 4 and line[0].isupper():
                return line
        return "Unknown"

    def _extract_email(self, text: str) -> str:
        m = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
        return m.group(0) if m else "N/A"

    def _extract_phone(self, text: str) -> str:
        m = re.search(r"(\+?\d{1,3}[\-.\s]?)?\(?\d{3}\)?[\-.\s]?\d{3}[\-.\s]?\d{4}", text)
        return m.group(0) if m else "N/A"

    def _extract_experience_years(self, text: str) -> int:
        matches = re.findall(r"(\d+)\s*\+?\s*years?", text.lower())
        return max((int(m) for m in matches), default=0)

    def _extract_skills_flat(self, text: str) -> List[str]:
        text_lower = text.lower()
        found = set()
        for category_skills in self.skill_keywords.values():
            for skill in category_skills:
                if skill.lower() in text_lower:
                    found.add(skill.title())
        return list(found)

    def _extract_certifications(self, text: str) -> List[str]:
        certs = []
        patterns = [
            r"(AWS\s+Certified[^\n]+)",
            r"(Google\s+Certified[^\n]+)",
            r"(Microsoft\s+Certified[^\n]+)",
            r"(Certified\s+[A-Z][^\n]{5,60})",
            r"(NPTEL[^\n]+)",
            r"(Coursera[^\n]+)",
            r"(Udemy[^\n]+)",
        ]
        for pat in patterns:
            for m in re.finditer(pat, text, re.IGNORECASE):
                cert = m.group(1).strip()
                if len(cert) > 5:
                    certs.append(cert[:120])
        return list(dict.fromkeys(certs))  # deduplicate preserving order

    def _extract_education(self, text: str) -> List[str]:
        education = []
        pattern = r"(B\.?S\.?|B\.?Tech\.?|M\.?S\.?|M\.?Tech\.?|Ph\.?D\.?|MBA|B\.?E\.?)\s+(?:in\s+)?([^,\n]{5,60})"
        for degree, field in re.findall(pattern, text, re.IGNORECASE):
            education.append(f"{degree} in {field.strip()}")
        return education or ["Not specified"]

    def _extract_projects_dict(self, text: str) -> Dict:
        """Extract projects into a dict {name: {description, skills, tools, outcome}}."""
        projects = {}
        section = re.search(
            r"(?:projects?|portfolio|personal\s+projects?)(.*?)(?:skills|education|experience|certifications|$)",
            text,
            re.IGNORECASE | re.DOTALL,
        )
        if not section:
            return projects

        proj_text = section.group(1)
        items = re.split(r"\n(?=\s*[-•*]|\s*[A-Z][A-Za-z\s]+\s*[|\-–])", proj_text)
        skill_words = set()
        for lst in self.skill_keywords.values():
            skill_words.update(lst)

        for item in items[:6]:
            item = item.strip()
            if len(item) < 20:
                continue
            title_match = re.match(r"[-•*]?\s*([A-Za-z][A-Za-z\s]+?)(?:\s*[|\-–]|\n|$)", item)
            if not title_match:
                continue
            name = title_match.group(1).strip()[:60]
            if len(name) < 3:
                continue
            # Extract skills mentioned in this project block
            block_lower = item.lower()
            proj_skills = [s.title() for s in skill_words if s in block_lower]
            projects[name] = {
                "description": item[:200],
                "skills": proj_skills,
                "tools": [],
                "outcome": "",
            }
        return projects

    # ─── Knowledge Graph ───────────────────────────────────────────────────────

    def _build_knowledge_graph(self, structured: Dict) -> Dict:
        """
        Build a project-skill relational knowledge graph.
        Example:
          {
            "projects": {
              "TalentScout": {"skills": ["Gemini API", "Streamlit", "SQLite"]},
              "Hygiene Monitor": {"skills": ["Random Forest", "Firebase", "IoT"]}
            },
            "experience": [
              {"company": "XYZ", "role": "Intern", "skills": ["Python", "FastAPI"]}
            ],
            "all_allowed_tech": ["gemini api", "streamlit", "sqlite", ...]
          }
        """
        kg_projects = {}
        for name, details in (structured.get("projects") or {}).items():
            all_tech = list(dict.fromkeys(
                (details.get("skills") or []) + (details.get("tools") or [])
            ))
            kg_projects[name] = {"skills": all_tech}

        kg_experience = []
        for exp in (structured.get("experience") or []):
            if isinstance(exp, dict):
                kg_experience.append({
                    "company": exp.get("company", ""),
                    "role": exp.get("role", ""),
                    "skills": exp.get("skills", []),
                })

        # Build the master allowed-tech set (lowercase) for grounding validation
        all_allowed = set()
        for s in (structured.get("skills") or []):
            all_allowed.add(s.lower())
        for t in (structured.get("tools") or []):
            all_allowed.add(t.lower())
        for details in kg_projects.values():
            for tech in details.get("skills", []):
                all_allowed.add(tech.lower())
        for exp in kg_experience:
            for tech in exp.get("skills", []):
                all_allowed.add(tech.lower())

        return {
            "projects": kg_projects,
            "experience": kg_experience,
            "all_allowed_tech": sorted(all_allowed),
        }

    # ─── Domain Detection ──────────────────────────────────────────────────────

    def _detect_domain(self, text: str) -> str:
        text_lower = text.lower()
        scores = {
            domain: sum(1 for kw in kws if kw in text_lower)
            for domain, kws in self.domain_keywords.items()
        }
        best = max(scores, key=scores.get)
        return best if scores[best] > 0 else "General"

    # ─── File Reading ──────────────────────────────────────────────────────────

    def _read_file(self, file_content: bytes, filename: str) -> str:
        fn = filename.lower()
        if fn.endswith(".pdf"):
            return self._extract_pdf_text(file_content)
        elif fn.endswith(".docx"):
            return self._extract_docx_text(file_content)
        elif fn.endswith(".txt"):
            return file_content.decode("utf-8")
        else:
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                return file_content.decode("latin-1")

    def _extract_pdf_text(self, file_content: bytes) -> str:
        import io
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            return "".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

    def _extract_docx_text(self, file_content: bytes) -> str:
        import io
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise